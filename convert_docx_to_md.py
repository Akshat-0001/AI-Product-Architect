import os
from docx import Document

def docx_to_markdown(docx_path, md_path):
    """Convert a DOCX file to Markdown format."""
    doc = Document(docx_path)
    markdown_content = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            markdown_content.append("")
            continue
        
        # Check for headings based on style
        style_name = paragraph.style.name if paragraph.style else "Normal"
        
        if style_name.startswith("Heading 1"):
            markdown_content.append(f"# {text}")
        elif style_name.startswith("Heading 2"):
            markdown_content.append(f"## {text}")
        elif style_name.startswith("Heading 3"):
            markdown_content.append(f"### {text}")
        elif style_name.startswith("Heading 4"):
            markdown_content.append(f"#### {text}")
        elif style_name.startswith("Heading 5"):
            markdown_content.append(f"##### {text}")
        elif style_name.startswith("Heading 6"):
            markdown_content.append(f"###### {text}")
        else:
            # Regular paragraph
            markdown_content.append(text)
    
    # Handle tables
    for table in doc.tables:
        markdown_content.append("")
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            markdown_content.append(" | ".join(row_data))
        markdown_content.append("")
    
    # Write to markdown file
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(markdown_content))
    
    print(f"Converted {docx_path} to {md_path}")

def main():
    docs_dir = r"c:\Users\comak\OneDrive\Desktop\AI Product Architect V2\docs"
    
    for filename in os.listdir(docs_dir):
        if filename.endswith('.docx'):
            docx_path = os.path.join(docs_dir, filename)
            md_filename = filename.replace('.docx', '.md')
            md_path = os.path.join(docs_dir, md_filename)
            docx_to_markdown(docx_path, md_path)

if __name__ == "__main__":
    main()
